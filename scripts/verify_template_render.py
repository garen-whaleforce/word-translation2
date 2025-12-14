#!/usr/bin/env python3
"""
驗證模板 render 結果
使用測試資料填充模板，檢查是否有殘留的 {{...}}
"""
import os
import re
import sys
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent.parent
TEMPLATE_PATH = ROOT / "templates" / "AST-B-MC-601.placeholder.v2.docx"
OUTPUT_PATH = ROOT / "templates" / "TEST_RENDER_OUTPUT.docx"

# 測試資料 - 模擬人工最終版的值
TEST_DATA = {
    # 封面
    "bsmi_designated_report_no": "SL2INT0157250509",
    "report_no": "AST-B-25120522-000",
    "applicant_name": "鼎福科技有限公司",
    "applicant_address": "新北市中和區民治街19巷8號",
    "cns_standard": "CNS 15598-1 (109年版)",
    "test_type": "型式試驗",
    "product_name_zh": "交流電源供應器",
    "main_model": "MC-601",
    "series_model": "A1231-200300C-US1",
    "trademark": "",
    "rated_input": "100-240 V～, 50/60 Hz, 1.7 A",
    "rated_output": "5.0 V  3.0 A 15.0 W or 9.0 V  3.0 A 27.0 W or\n     15.0 V  3.0 A 45.0 W or 20.0 V  3.0 A 60.0 W or\n     5.0 V-20.0 V  3.0 A 60.0 W(PPS)",
    "not_applicable_items": "不適用",
    "sample_conforms": "符合",
    "sample_not_conforms": "不符合",
    "sample_received_date": "114 年 12 月05 日",
    "test_date": "114 年 07 月28 日",
    "issue_date": "114 年 07 月29 日",
    "lab_name": "安捷檢測有限公司",
    "lab_address": "新北市新店區寶興路45巷8弄16號4樓",
    "overall_result": "符合",
    "report_author": "測試製作者",
    "report_signer": "測試簽署人",

    # 修訂紀錄
    "rev1_item": "01",
    "rev1_date": "114.07.29",
    "rev1_report_no": "AST-B-25120522-000",
    "rev1_desc": "主報告",

    # 試驗樣品特性
    "supply_connection_type": "A 型插接式設備",
    "protective_device_rated_current": "20 A建築",
    "equipment_mobility": "直插式設備",
    "ovc": "OVC II",
    "protection_class": "Class II",
    "special_installation": "無特殊安裝",
    "tma_c": "45 °C",
    "ip_rating": "IPX0",
    "equipment_altitude": "2000 m 或更低",
    "lab_altitude": "2000 m 或更低",
    "eut_mass_kg": "0.072 kg",
}


def replace_placeholders_in_cell(cell, data: dict) -> list:
    """替換儲存格中的 placeholders，回傳替換紀錄"""
    replacements = []

    for para in cell.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = original_text

            # 找出所有 {{...}} 並替換
            matches = re.findall(r'\{\{([a-z_0-9]+)\}\}', new_text)
            for placeholder in matches:
                if placeholder in data:
                    value = data[placeholder]
                    new_text = new_text.replace(f"{{{{{placeholder}}}}}", value)
                    replacements.append((placeholder, value[:30] if value else "(空)"))

            if new_text != original_text:
                run.text = new_text

    return replacements


def render_template(template_path: Path, output_path: Path, data: dict):
    """填充模板並儲存"""
    print(f"讀取模板: {template_path}")
    doc = Document(str(template_path))

    all_replacements = []

    # 處理所有表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                replacements = replace_placeholders_in_cell(cell, data)
                for placeholder, value in replacements:
                    all_replacements.append(f"Table {ti}, Row {ri}, Col {ci}: {{{{{placeholder}}}}} -> {value}")

    # 處理文件主體段落
    for para in doc.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = original_text
            matches = re.findall(r'\{\{([a-z_0-9]+)\}\}', new_text)
            for placeholder in matches:
                if placeholder in data:
                    value = data[placeholder]
                    new_text = new_text.replace(f"{{{{{placeholder}}}}}", value)
                    all_replacements.append(f"Paragraph: {{{{{placeholder}}}}} -> {value[:30]}")
            if new_text != original_text:
                run.text = new_text

    print(f"\n共替換 {len(all_replacements)} 處")

    doc.save(str(output_path))
    print(f"儲存至: {output_path}")

    return doc


def check_remaining_placeholders(doc) -> list:
    """檢查是否有殘留的 {{...}}"""
    remaining = []

    # 檢查表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text
                matches = re.findall(r'\{\{([a-z_0-9]+)\}\}', text)
                for m in matches:
                    remaining.append(f"Table {ti}, Row {ri}, Col {ci}: {{{{{m}}}}}")

    # 檢查段落
    for pi, para in enumerate(doc.paragraphs):
        text = para.text
        matches = re.findall(r'\{\{([a-z_0-9]+)\}\}', text)
        for m in matches:
            remaining.append(f"Paragraph {pi}: {{{{{m}}}}}")

    return remaining


def main():
    os.chdir(ROOT)

    # Render 模板
    doc = render_template(TEMPLATE_PATH, OUTPUT_PATH, TEST_DATA)

    # 檢查殘留
    print("\n=== 檢查殘留的 placeholders ===")
    remaining = check_remaining_placeholders(doc)

    if remaining:
        print(f"⚠️ 發現 {len(remaining)} 個未替換的 placeholder:")
        for r in remaining:
            print(f"  - {r}")
    else:
        print("✅ 沒有殘留的 placeholder")

    # 檢查缺值欄位
    print("\n=== 檢查缺值欄位 ===")
    missing_values = [k for k, v in TEST_DATA.items() if not v]
    if missing_values:
        print(f"⚠️ 以下欄位為空值:")
        for m in missing_values:
            print(f"  - {{{{{m}}}}}")
    else:
        print("✅ 所有欄位都有值")

    print(f"\n📄 測試輸出檔案: {OUTPUT_PATH}")
    print("請用 Word 開啟檢查結果是否與人工最終版一致")


if __name__ == "__main__":
    main()
