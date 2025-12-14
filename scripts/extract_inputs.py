# scripts/extract_inputs.py
from pathlib import Path
import json
import re

DOCX_PATH = Path("templates/AST-B-MC-601.docx")
PDF_TEXT_PATH = Path("artifacts/cb_mc_601_text.txt")  # 先用 adobe/azure 轉成 text 放這裡
OUT_DIR = Path("artifacts")
OUT_DIR.mkdir(parents=True, exist_ok=True)

def extract_docx_text(docx_path: Path):
    from docx import Document
    doc = Document(str(docx_path))
    blocks = []
    # paragraphs
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            blocks.append({"type": "p", "idx": i, "text": t[:500]})
    # tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = (cell.text or "").strip()
                if t:
                    blocks.append({"type": "cell", "table": ti, "row": ri, "col": ci, "text": t[:500]})
    return blocks

def extract_candidate_fields(blocks):
    """
    找出疑似「可回填欄位」：包含冒號、括號欄位、固定關鍵詞（型號/額定/報告號/證書號等）
    產生短清單供 LLM 對照 CB text 做 mapping
    """
    keywords = [
        "報告", "證書", "型號", "系列型號", "額定", "輸入", "輸出", "製造", "生產廠", "地址",
        "標準", "IEC", "IP", "海拔", "重量", "質量", "功率", "負載"
    ]
    cand = []
    for b in blocks:
        txt = b["text"]
        if any(k in txt for k in keywords) or ":" in txt or "：" in txt:
            cand.append(b)
    return cand[:800]  # 控制大小，避免又爆

def load_pdf_text(path: Path):
    if not path.exists():
        raise SystemExit(f"Missing {path}. Please convert PDF to text first (adobe/azure) and save it.")
    t = path.read_text(encoding="utf-8", errors="ignore")
    t = re.sub(r"[ \t]+", " ", t)
    # 只保留前後與關鍵段落（避免太長）
    lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
    return lines

def compress_pdf_lines(lines):
    keep = []
    key_pat = re.compile(r"(Report Number|Certificate|Standard|Model|Type|Ratings|Input|Output|Tma|IP|altitude|mass|factory|manufacturer)", re.I)
    for ln in lines:
        if key_pat.search(ln):
            keep.append(ln)
    # 補一些上下文（前後各 1 行）
    enriched = []
    s = set()
    for i, ln in enumerate(lines):
        if ln in keep:
            for j in (i-1, i, i+1):
                if 0 <= j < len(lines):
                    if j not in s:
                        enriched.append(lines[j])
                        s.add(j)
    return enriched[:1200]

def main():
    blocks = extract_docx_text(DOCX_PATH)
    cand = extract_candidate_fields(blocks)
    pdf_lines = compress_pdf_lines(load_pdf_text(PDF_TEXT_PATH))

    (OUT_DIR / "ast_blocks.json").write_text(json.dumps(blocks, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT_DIR / "ast_candidates.json").write_text(json.dumps(cand, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT_DIR / "cb_keylines.txt").write_text("\n".join(pdf_lines), encoding="utf-8")

    print("Wrote:")
    print("- artifacts/ast_blocks.json")
    print("- artifacts/ast_candidates.json")
    print("- artifacts/cb_keylines.txt")

if __name__ == "__main__":
    main()
