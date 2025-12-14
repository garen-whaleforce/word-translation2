#!/usr/bin/env python3
"""
更新 AST-B-MC-601.placeholder.docx 模板
在指定位置插入 {{placeholder}} 格式的佔位符
"""
import os
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt

# 專案根目錄
ROOT = Path(__file__).parent.parent
TEMPLATE_DIR = ROOT / "templates"
INPUT_PATH = TEMPLATE_DIR / "AST-B-MC-601.placeholder.docx"
OUTPUT_PATH = TEMPLATE_DIR / "AST-B-MC-601.placeholder.v2.docx"


def set_cell_text(cell, text: str, preserve_first_run: bool = True):
    """
    設定儲存格文字，確保 placeholder 是連續字串（單一 run）

    Args:
        cell: Word 表格儲存格
        text: 要設定的文字
        preserve_first_run: 是否保留第一個 run 的格式
    """
    if not cell.paragraphs:
        cell.text = text
        return

    para = cell.paragraphs[0]

    if preserve_first_run and para.runs:
        # 保留第一個 run 的格式，清空其他 run
        first_run = para.runs[0]
        first_run.text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        # 直接設定段落文字
        para.clear()
        para.add_run(text)


def get_cell_text(cell) -> str:
    """取得儲存格的完整文字"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def update_template():
    """更新模板，插入 placeholders"""
    print(f"讀取模板: {INPUT_PATH}")
    doc = Document(str(INPUT_PATH))

    changes = []

    # ========================================
    # Table 0: 封面
    # ========================================
    table0 = doc.tables[0]

    # 定義封面欄位對照表: (row, col, placeholder, is_fixed_value)
    cover_fields = [
        # Row 2: 標準檢驗局試驗報告指定編號 (值在 col 2)
        (2, 2, "{{bsmi_designated_report_no}}", False),

        # Row 3: 報告編號
        (3, 2, "{{report_no}}", False),

        # Row 4: 申請者名稱
        (4, 2, "{{applicant_name}}", False),

        # Row 5: 申請者地址
        (5, 2, "{{applicant_address}}", False),

        # Row 6: 生產廠場 - 固定值
        (6, 2, "詳見報告第4頁", True),

        # Row 7: 生產廠場地址 - 固定值
        (7, 2, "詳見報告第4頁", True),

        # Row 8: 試驗標準(規範)
        (8, 2, "{{cns_standard}}", False),

        # Row 9: 試驗方式
        (9, 2, "{{test_type}}", False),

        # Row 10: 品名
        (10, 2, "{{product_name_zh}}", False),

        # Row 11: 主型號
        (11, 2, "{{main_model}}", False),

        # Row 12: 系列型號
        (12, 2, "{{series_model}}", False),

        # Row 13: 廠牌/商標 - 已有 {{trademark}}，保持不變
        # (13, 2, "{{trademark}}", False),  # 已存在

        # Row 14: 額定 - 特殊處理，需要輸入和輸出
        (14, 2, "輸入: {{rated_input}}\n輸出: {{rated_output}}", False),

        # Row 15: 測試項目不適用
        (15, 2, "{{not_applicable_items}}", False),

        # Row 16: 測試樣品符合要求
        (16, 2, "{{sample_conforms}}", False),

        # Row 17: 測試樣品不符合要求
        (17, 2, "{{sample_not_conforms}}", False),

        # Row 18: 試驗件收件日
        (18, 2, "{{sample_received_date}}", False),

        # Row 19: 執行測試日
        (19, 2, "{{test_date}}", False),

        # Row 20: 報告發行日
        (20, 2, "{{issue_date}}", False),

        # Row 21: 測試單位
        (21, 2, "{{lab_name}}", False),

        # Row 22: 測試單位地址
        (22, 2, "{{lab_address}}", False),

        # Row 23: 試驗結果
        (23, 2, "{{overall_result}}", False),

        # Row 25: 報告製作者/簽署人 - 已有 placeholder，保持不變
    ]

    print("\n=== 更新封面 (Table 0) ===")
    for row_idx, col_idx, new_text, is_fixed in cover_fields:
        if row_idx < len(table0.rows):
            row = table0.rows[row_idx]
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                old_text = get_cell_text(cell)[:50]
                set_cell_text(cell, new_text)

                # 由於合併儲存格，col 2-4 可能是同一個 cell
                # 確保 col 3 和 col 4 也被設定（如果它們是不同的 cell）
                for extra_col in [3, 4]:
                    if extra_col < len(row.cells):
                        extra_cell = row.cells[extra_col]
                        # 檢查是否為同一個 cell（合併儲存格）
                        if extra_cell._tc != cell._tc:
                            set_cell_text(extra_cell, new_text)

                status = "固定值" if is_fixed else "placeholder"
                changes.append(f"[封面 Row {row_idx}] {status}: {new_text[:40]}...")
                print(f"  Row {row_idx}: '{old_text}' -> '{new_text[:40]}...'")

    # ========================================
    # Table 2: 報告修訂紀錄
    # ========================================
    table2 = doc.tables[2]

    print("\n=== 更新報告修訂紀錄 (Table 2) ===")

    # Row 2 是第一筆資料列 (Row 0 是標題, Row 1 是欄位名稱)
    if len(table2.rows) > 2:
        rev_row = table2.rows[2]
        rev_fields = [
            (0, "{{rev1_item}}"),
            (1, "{{rev1_date}}"),
            (2, "{{rev1_report_no}}"),
            (3, "{{rev1_desc}}"),
        ]

        for col_idx, placeholder in rev_fields:
            if col_idx < len(rev_row.cells):
                cell = rev_row.cells[col_idx]
                old_text = get_cell_text(cell)[:30]
                set_cell_text(cell, placeholder)
                changes.append(f"[修訂紀錄 Col {col_idx}] {placeholder}")
                print(f"  Col {col_idx}: '{old_text}' -> '{placeholder}'")

    # ========================================
    # Table 3: 試驗樣品特性
    # ========================================
    table3 = doc.tables[3]

    print("\n=== 更新試驗樣品特性 (Table 3) ===")

    # 試驗樣品特性欄位對照表: (row, col, placeholder)
    # 注意：這些欄位的原始格式有 checkbox 符號，我們只在值的位置加入 placeholder
    particulars_fields = [
        # Row 5: 電源連接方式 (值在 col 1 的開頭位置)
        (5, 1, "{{supply_connection_type}}"),

        # Row 6: 考慮保護裝置的額定電流
        (6, 1, "{{protective_device_rated_current}}"),

        # Row 7: 設備移動性
        (7, 1, "{{equipment_mobility}}"),

        # Row 8: 過壓類別(OVC)
        (8, 1, "{{ovc}}"),

        # Row 9: 防電擊保護 (Class)
        (9, 1, "{{protection_class}}"),

        # Row 10: 特殊安裝位置
        (10, 1, "{{special_installation}}"),

        # Row 11: 污染等級 - 保留原有格式，加入 placeholder
        # (11, 1, "{{pollution_degree}}"),  # 格式較複雜，暫不處理

        # Row 12: 製造商宣告Tma
        (12, 1, "{{tma_c}}"),

        # Row 13: IP等級
        (13, 1, "{{ip_rating}}"),

        # Row 15: 設備適用的海拔高度
        (15, 1, "{{equipment_altitude}}"),

        # Row 16: 測試實驗室海拔高度
        (16, 1, "{{lab_altitude}}"),

        # Row 17: 設備質量(kg)
        (17, 1, "{{eut_mass_kg}}"),
    ]

    for row_idx, col_idx, placeholder in particulars_fields:
        if row_idx < len(table3.rows):
            row = table3.rows[row_idx]
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                old_text = get_cell_text(cell)[:40]
                set_cell_text(cell, placeholder)
                changes.append(f"[試驗樣品 Row {row_idx}] {placeholder}")
                print(f"  Row {row_idx}: '{old_text}' -> '{placeholder}'")

    # ========================================
    # 儲存檔案
    # ========================================
    print(f"\n儲存到: {OUTPUT_PATH}")
    doc.save(str(OUTPUT_PATH))

    print(f"\n=== 完成！共更新 {len(changes)} 處 ===")
    for c in changes:
        print(f"  - {c}")

    return OUTPUT_PATH


def verify_placeholders(doc_path: Path):
    """驗證模板中的 placeholders"""
    print(f"\n=== 驗證 placeholders ===")
    doc = Document(str(doc_path))

    found_placeholders = set()

    # 掃描所有表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text
                # 找出所有 {{...}} 格式的 placeholder
                import re
                matches = re.findall(r'\{\{([a-z_0-9]+)\}\}', text)
                for m in matches:
                    found_placeholders.add(m)

    print(f"找到 {len(found_placeholders)} 個 placeholders:")
    for p in sorted(found_placeholders):
        print(f"  - {{{{{p}}}}}")

    # 檢查預期的 placeholders
    expected = {
        # 封面
        "bsmi_designated_report_no", "report_no", "applicant_name", "applicant_address",
        "cns_standard", "test_type", "product_name_zh", "main_model", "series_model",
        "trademark", "rated_input", "rated_output", "not_applicable_items",
        "sample_conforms", "sample_not_conforms", "sample_received_date",
        "test_date", "issue_date", "lab_name", "lab_address", "overall_result",
        "report_author", "report_signer",
        # 修訂紀錄
        "rev1_item", "rev1_date", "rev1_report_no", "rev1_desc",
        # 試驗樣品特性
        "supply_connection_type", "protective_device_rated_current",
        "equipment_mobility", "ovc", "protection_class", "special_installation",
        "tma_c", "ip_rating", "equipment_altitude", "lab_altitude", "eut_mass_kg",
    }

    missing = expected - found_placeholders
    extra = found_placeholders - expected

    if missing:
        print(f"\n⚠️ 缺少的 placeholders ({len(missing)} 個):")
        for p in sorted(missing):
            print(f"  - {{{{{p}}}}}")

    if extra:
        print(f"\n📋 額外的 placeholders ({len(extra)} 個):")
        for p in sorted(extra):
            print(f"  - {{{{{p}}}}}")

    if not missing:
        print("\n✅ 所有預期的 placeholders 都已存在")

    return found_placeholders, missing


if __name__ == "__main__":
    # 確保工作目錄正確
    os.chdir(ROOT)

    # 更新模板
    output_path = update_template()

    # 驗證結果
    verify_placeholders(output_path)

    print(f"\n📄 新模板已儲存至: {output_path}")
    print("請用 Word 開啟檢查格式是否正確")
