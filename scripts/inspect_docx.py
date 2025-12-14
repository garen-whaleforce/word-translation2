# scripts/inspect_docx.py
"""檢查 docx 表格結構"""
from pathlib import Path
from docx import Document

INPUT_PATH = Path("templates/AST-B-MC-601.docx")

doc = Document(str(INPUT_PATH))

print(f"共 {len(doc.tables)} 個表格\n")

for ti, table in enumerate(doc.tables):
    print(f"=== Table {ti} ({len(table.rows)} rows) ===")
    for ri, row in enumerate(table.rows):
        print(f"  Row {ri}:")
        for ci, cell in enumerate(row.cells):
            text = cell.text.replace("\n", "\\n")[:60]
            print(f"    Col {ci}: {text}")
    print()
    if ti >= 4:  # 只顯示前 5 個表格
        print("... (省略其餘表格)")
        break
