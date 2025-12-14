# scripts/create_placeholder_template.py
"""
將 AST-B-MC-601.docx 中的特定文字替換為 placeholder
只替換文字內容，不改變任何樣式
"""
from pathlib import Path
from docx import Document
import re

INPUT_PATH = Path("templates/AST-B-MC-601.docx")
OUTPUT_PATH = Path("templates/AST-B-MC-601.placeholder.docx")

# Placeholder 對照表：(table_idx, row_idx, col_idx) -> (placeholder_name, template_text)
# template_text 為 None 時自動生成 {{placeholder_name}}
CELL_PLACEHOLDERS = {
    # 封面表格 (table 0) - 值在 col 2
    (0, 2, 2): ("bsmi_report_number", None),      # 標準檢驗局試驗報告指定編號
    (0, 3, 2): ("report_number", None),           # 報告編號
    (0, 4, 2): ("applicant_name", None),          # 申請者名稱
    (0, 5, 2): ("applicant_address", None),       # 地址
    (0, 6, 2): ("factory_name", "{{#factories}}{{name}}{{/factories}}"),  # 生產廠場
    (0, 7, 2): ("factory_address", "{{#factories}}{{address}}{{/factories}}"),  # 工廠地址
    (0, 8, 2): ("test_standard", None),           # 試驗標準
    (0, 9, 2): ("test_method", None),             # 試驗方式
    (0, 10, 2): ("product_name", None),           # 品名
    (0, 11, 2): ("main_model", None),             # 主型號
    (0, 12, 2): ("series_models", "{{#series_models}}{{.}}{{/series_models}}"),  # 系列型號
    (0, 13, 2): ("trademark", None),              # 廠牌/商標
    (0, 14, 2): ("ratings", "輸入: {{ratings.input}}\n輸出: {{ratings.output}}"),  # 額定
    (0, 15, 2): ("test_na_items", None),          # 測試項目不適用
    (0, 16, 2): ("test_pass_items", None),        # 測試樣品符合要求
    (0, 17, 2): ("test_fail_items", None),        # 測試樣品不符合要求
    (0, 18, 2): ("sample_receipt_date", None),    # 試驗件收件日
    (0, 19, 2): ("test_date", None),              # 執行測試日
    (0, 20, 2): ("report_issue_date", None),      # 報告發行日
    (0, 21, 2): ("testing_lab", None),            # 測試單位
    (0, 22, 2): ("testing_lab_address", None),    # 測試單位地址
    (0, 23, 2): ("test_result", None),            # 試驗結果
    (0, 25, 1): ("report_author", None),          # 報告製作者
    (0, 25, 4): ("report_signer", None),          # 報告簽署人

    # 修訂紀錄表格 (table 2) - 資料行 (row 2 為第一筆資料)
    (2, 2, 0): ("revision_idx", "{{#revision_history}}{{@index}}{{/revision_history}}"),
    (2, 2, 1): ("revision_date", "{{#revision_history}}{{date}}{{/revision_history}}"),
    (2, 2, 2): ("revision_report_number", "{{#revision_history}}{{report_number}}{{/revision_history}}"),
    (2, 2, 3): ("revision_description", "{{#revision_history}}{{description}}{{/revision_history}}"),

    # 試驗樣品特性表格 (table 3)
    (3, 5, 1): ("supply_connection_type", None),   # 電源連接方式
    (3, 6, 1): ("protection_device_rating", None), # 保護裝置額定電流
    (3, 7, 1): ("equipment_mobility", None),       # 設備移動性
    (3, 9, 1): ("equipment_class", None),          # 防電擊保護 (Class)
    (3, 10, 1): ("special_installation", None),    # 特殊安裝位置
    (3, 12, 1): ("tma", None),                     # Tma
    (3, 13, 1): ("ip_rating", None),               # IP 等級
    (3, 15, 1): ("altitude", None),                # 海拔高度
    (3, 16, 1): ("test_lab_altitude", None),       # 測試實驗室海拔
    (3, 17, 1): ("equipment_mass", None),          # 設備質量
}


def replace_text_preserve_style(cell, new_text):
    """
    替換 cell 中的文字，保留原有樣式
    策略：找到第一個有文字的 run，替換其文字，清空其他 run
    """
    for para in cell.paragraphs:
        if not para.runs:
            continue

        # 找到第一個非空 run
        first_run_with_text = None
        for run in para.runs:
            if run.text.strip():
                first_run_with_text = run
                break

        if first_run_with_text:
            # 替換第一個 run 的文字
            first_run_with_text.text = new_text
            # 清空其他 run
            for run in para.runs:
                if run != first_run_with_text:
                    run.text = ""
            return True

    # 如果沒有 run，直接設定段落文字
    if cell.paragraphs:
        cell.paragraphs[0].text = new_text
        return True

    return False


def get_cell_text(cell):
    """取得 cell 的完整文字"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def main():
    doc = Document(str(INPUT_PATH))

    replaced_count = 0
    replaced_list = []
    not_found = []

    # 處理表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                key = (table_idx, row_idx, col_idx)

                # 檢查 CELL_PLACEHOLDERS
                if key in CELL_PLACEHOLDERS:
                    placeholder_name, template_text = CELL_PLACEHOLDERS[key]
                    old_text = get_cell_text(cell)

                    # 決定替換文字
                    if template_text is not None:
                        new_text = template_text
                    else:
                        new_text = "{{" + placeholder_name + "}}"

                    if replace_text_preserve_style(cell, new_text):
                        replaced_count += 1
                        replaced_list.append(f"{placeholder_name} @ table:{table_idx}, row:{row_idx}, col:{col_idx}")
                        old_preview = old_text[:40].replace('\n', '\\n')
                        print(f"[OK] {placeholder_name}: '{old_preview}...' -> '{new_text[:50]}'")
                    else:
                        not_found.append(f"{placeholder_name} @ table:{table_idx}, row:{row_idx}, col:{col_idx}")

    # 儲存
    doc.save(str(OUTPUT_PATH))

    print("\n" + "="*60)
    print(f"完成！共替換 {replaced_count} 處")
    print(f"輸出檔案：{OUTPUT_PATH}")
    print(f"\n替換的欄位列表：")
    for item in replaced_list:
        print(f"  - {item}")

    if not_found:
        print(f"\n無法定位的 placeholder ({len(not_found)} 項)：")
        for item in not_found:
            print(f"  - {item}")

    # 列出尚未處理的 placeholder（在 schema 中但未在此腳本定義）
    all_schema_placeholders = [
        "cb_certificate_number", "cb_report_number",
        "limiting_component_voltage", "insulation_resistance", "multiplication_factor",
        "outdoor_min_temp", "bsmi_lab_code"
    ]

    defined_names = {v[0] for v in CELL_PLACEHOLDERS.values()}
    undefined = [p for p in all_schema_placeholders if p not in defined_names]

    if undefined:
        print(f"\n尚未定位的 placeholder（需在其他表格中確認位置）：")
        for p in undefined:
            print(f"  - {p}")


if __name__ == "__main__":
    main()
