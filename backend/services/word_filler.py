"""
==============================================
Word Template Filler Service
使用 python-docx 填寫 CNS 報告 Word 模板
==============================================

此模組負責：
1. 讀取 Word 模板
2. 尋找並替換 placeholder（{{...}}）
3. 處理 checkbox（□ → ■）
4. 處理 Word FORMCHECKBOX 控制項的勾選狀態
5. 填寫表格中的系列型號
6. 確保不破壞原有格式

⚠️ 重要設計原則：
- 只修改文字內容（w:t），不新增/刪除段落或表格
- 處理 Word 將 placeholder 切成多個 run 的情況
- 保持原有的字體、大小、顏色等格式
- FORMCHECKBOX 需要直接操作 XML
"""

import re
import os
import glob
import zipfile
import shutil
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from copy import deepcopy
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from schemas.report_schema import ReportSchema, BasicInfo
from config import settings
from utils.logger import get_logger

logger = get_logger(__name__)


# ==============================================
# Constants
# ==============================================

# Placeholder 正則表達式
PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
PLACEHOLDER_REGEX_RAW = re.compile(r'\{\{.*?\}\}')

# Checkbox 符號
CHECKBOX_UNCHECKED = "□"
CHECKBOX_CHECKED = "■"

# 預設最大系列型號數量（對應 Word 模板中的 {{series_model_N}}）
MAX_SERIES_MODELS = 60


# ==============================================
# Placeholder Mapping Builder
# ==============================================

def build_placeholder_mapping(schema: ReportSchema) -> Dict[str, str]:
    """
    根據 ReportSchema 建立 placeholder 到值的對應表

    Args:
        schema: ReportSchema 物件

    Returns:
        Dict[placeholder_name, value]
    """
    mapping = {}

    # ===================
    # Basic Info
    # ===================
    bi = schema.basic_info
    ast_report_no = bi.ast_report_no or bi.cb_report_no or ""
    mapping["report_no"] = bi.cb_report_no or ""
    mapping["cb_report_no"] = bi.cb_report_no or ""
    mapping["ast_report_no"] = ast_report_no
    mapping["header_report_no"] = ast_report_no
    mapping["cns_report_no"] = bi.cns_report_no or ""
    mapping["standard"] = bi.standard or ""
    mapping["standard_version"] = bi.standard_version or ""
    mapping["national_differences"] = bi.national_differences or ""
    mapping["test_lab"] = bi.test_lab or ""
    mapping["test_lab_country"] = bi.test_lab_country or ""

    mapping["applicant_en"] = bi.applicant_en or ""
    mapping["applicant_address_en"] = bi.applicant_address_en or ""
    mapping["manufacturer_en"] = bi.manufacturer_en or ""
    mapping["manufacturer_address_en"] = bi.manufacturer_address_en or ""
    mapping["factory_name_en"] = bi.factory_name_en or ""
    mapping["factory_address_en"] = bi.factory_address_en or ""

    mapping["product_name_en"] = bi.product_name_en or ""
    mapping["model_main"] = bi.model_main or ""
    mapping["main_model"] = bi.model_main or ""  # alias for model_main
    mapping["brand"] = bi.brand or ""
    mapping["trademark"] = bi.trademark or ""

    mapping["ratings_input"] = bi.ratings_input or ""
    mapping["rated_input"] = bi.ratings_input or ""  # alias for ratings_input
    mapping["ratings_output"] = bi.ratings_output or ""
    mapping["rated_output"] = bi.ratings_output or ""  # alias for ratings_output
    if bi.rated_output_lines:
        mapping["rated_output_block"] = "\n".join(bi.rated_output_lines)
    else:
        mapping["rated_output_block"] = bi.ratings_output or ""
    mapping["ratings_power"] = bi.ratings_power or ""
    # 最大輸出（若未提供，嘗試從 rated_output_lines 推估）
    mapping["max_output_w"] = bi.max_output_w or ""
    mapping["max_output_v"] = bi.max_output_v or ""
    mapping["max_output_a"] = bi.max_output_a or ""
    if not mapping["max_output_w"]:
        numbers = []
        for line in (bi.rated_output_lines or []):
            match = re.findall(r"([0-9]+(?:\.[0-9]+)?)\s*W", line, flags=re.IGNORECASE)
            numbers.extend(match)
        if numbers:
            mapping["max_output_w"] = str(max(float(n) for n in numbers))

    mapping["issue_date"] = bi.issue_date or ""
    mapping["issue_date_short"] = bi.issue_date_short or ""
    mapping["receive_date"] = bi.receive_date or ""
    mapping["sample_received_date"] = bi.receive_date or ""  # alias for receive_date
    mapping["test_date_from"] = bi.test_date_from or ""
    mapping["test_date"] = bi.test_date_from or ""  # alias for test_date_from
    mapping["test_date_to"] = bi.test_date_to or ""

    # CB 報告資訊
    mapping["cb_test_lab"] = bi.cb_test_lab or ""
    mapping["cb_certificate_no"] = bi.cb_certificate_no or ""
    mapping["cb_standard"] = bi.cb_standard or ""

    # 設備資訊
    mapping["equipment_mass"] = bi.equipment_mass or ""
    mapping["eut_mass_kg"] = bi.equipment_mass or ""  # alias for equipment_mass
    mapping["protection_rating"] = bi.protection_rating or ""
    mapping["protective_device_rated_current"] = bi.protection_rating or ""  # alias
    mapping["national_differences_summary"] = bi.national_differences_summary or ""
    mapping["model_differences_block"] = bi.model_differences or ""
    mapping["cb_report_note"] = bi.cb_report_note or ""
    mapping["temperature_requirements_text"] = bi.temperature_requirements_text or ""

    # ===================
    # Translations (繁中)
    # ===================
    trans = schema.translations
    mapping["applicant_zh"] = trans.applicant_zh or ""
    mapping["applicant_name"] = trans.applicant_zh or ""  # alias for applicant_zh
    mapping["applicant_address_zh"] = trans.applicant_address_zh or ""
    mapping["applicant_address"] = trans.applicant_address_zh or ""  # alias
    mapping["manufacturer_zh"] = trans.manufacturer_zh or ""
    mapping["manufacturer_address_zh"] = trans.manufacturer_address_zh or ""
    mapping["product_name_zh"] = trans.product_name_zh or ""
    mapping["factory_name_zh"] = trans.factory_name_zh or ""
    mapping["factory_address_zh"] = trans.factory_address_zh or ""
    # 多工廠支援
    mapping["factory_name_1"] = trans.factory_name_1 or ""
    mapping["factory_address_1"] = trans.factory_address_1 or ""
    mapping["factory_name_2"] = trans.factory_name_2 or ""
    mapping["factory_address_2"] = trans.factory_address_2 or ""
    # 工廠清單（動態列表）
    if schema.factories:
        factory_lines = []
        for f in schema.factories:
            parts = [p for p in [f.name, f.address] if p]
            if parts:
                factory_lines.append(" / ".join(parts))
        mapping["factory_list"] = "; ".join(factory_lines)
    else:
        mapping["factory_list"] = ""

    # ===================
    # Test Item Particulars
    # ===================
    tip = schema.test_item_particulars
    mapping["product_group"] = tip.product_group or ""
    mapping["ovc"] = tip.ovc or ""
    mapping["pollution_degree"] = tip.pollution_degree or ""
    mapping["ip_code"] = tip.ip_code or ""
    mapping["ip_rating"] = tip.ip_code or ""  # alias for ip_code
    mapping["tma"] = tip.tma or ""
    mapping["tma_c"] = tip.tma or ""  # alias for tma
    mapping["altitude_limit_m"] = str(tip.altitude_limit_m) if tip.altitude_limit_m else ""
    mapping["altitude"] = f"{tip.altitude_limit_m} m 或更低" if tip.altitude_limit_m else ""
    mapping["equipment_altitude"] = str(tip.altitude_limit_m) if tip.altitude_limit_m else ""  # alias
    mapping["mains_supply"] = tip.mains_supply or ""
    mapping["rated_voltage"] = tip.rated_voltage or ""
    mapping["rated_frequency"] = tip.rated_frequency or ""
    mapping["rated_current"] = tip.rated_current or ""
    mapping["protection_class"] = tip.protection_class or ""
    mapping["insulation_type"] = tip.insulation_type or ""
    # 設備移動性：清空文字內容，改用 checkbox 顯示
    # （checkbox 已在 FORMCHECKBOX 處理邏輯中勾選）
    mapping["mobility"] = ""
    mapping["equipment_mobility"] = ""  # alias for mobility

    # Classification of use（多選，用逗號分隔）
    mapping["classification_of_use"] = ", ".join(tip.classification_of_use) if tip.classification_of_use else ""
    # Supply connection（多選）
    mapping["supply_connection"] = ", ".join(tip.supply_connection) if tip.supply_connection else ""
    mapping["supply_connection_type"] = ", ".join(tip.supply_connection) if tip.supply_connection else ""  # alias

    # ===================
    # Series Models
    # ===================
    # 單一 series_model 欄位（所有型號用逗號分隔）
    all_models = [m.model for m in schema.series_models if m.model]
    mapping["series_model"] = ", ".join(all_models) if all_models else ""
    mapping["model_list"] = mapping["series_model"]

    for i, model in enumerate(schema.series_models[:MAX_SERIES_MODELS], start=1):
        mapping[f"series_model_{i}"] = model.model or ""
        mapping[f"series_model_{i}_vout"] = model.vout or ""
        mapping[f"series_model_{i}_iout"] = model.iout or ""
        mapping[f"series_model_{i}_pout"] = model.pout or ""
        mapping[f"series_model_{i}_vin"] = model.vin or ""
        mapping[f"series_model_{i}_iin"] = model.iin or ""
        mapping[f"series_model_{i}_case_type"] = model.case_type or ""
        mapping[f"series_model_{i}_connector"] = model.connector_type or ""
        mapping[f"series_model_{i}_diff"] = model.differences or ""
        mapping[f"series_model_{i}_remarks"] = model.remarks or ""

    # 填充剩餘的系列型號為空字串（避免 {{series_model_N}} 殘留）
    for i in range(len(schema.series_models) + 1, MAX_SERIES_MODELS + 1):
        mapping[f"series_model_{i}"] = ""
        mapping[f"series_model_{i}_vout"] = ""
        mapping[f"series_model_{i}_iout"] = ""
        mapping[f"series_model_{i}_pout"] = ""
        mapping[f"series_model_{i}_vin"] = ""
        mapping[f"series_model_{i}_iin"] = ""
        mapping[f"series_model_{i}_case_type"] = ""
        mapping[f"series_model_{i}_connector"] = ""
        mapping[f"series_model_{i}_diff"] = ""
        mapping[f"series_model_{i}_remarks"] = ""

    # ===================
    # New Fields (新增欄位)
    # ===================
    # BSMI 相關
    mapping["bsmi_designated_report_no"] = bi.bsmi_designated_report_no or ""
    mapping["cns_standard"] = bi.cns_standard or ""
    mapping["cns_standard_version"] = bi.cns_standard_version or ""

    # 試驗結果相關
    mapping["test_type"] = bi.test_type or ""
    mapping["overall_result"] = bi.overall_result or ""
    mapping["sample_conforms"] = bi.sample_conforms or ""
    mapping["sample_not_conforms"] = bi.sample_not_conforms or ""
    mapping["not_applicable_items"] = bi.not_applicable_items or ""
    mapping["special_installation"] = bi.special_installation or ""

    # ===================
    # Revision Records (修訂記錄)
    # ===================
    # 預設最多支援 5 筆修訂記錄
    for i, rev in enumerate(schema.revision_records[:5], start=1):
        mapping[f"rev{i}_item"] = rev.item or ""
        mapping[f"rev{i}_date"] = rev.date or ""
        mapping[f"rev{i}_report_no"] = rev.report_no or ""
        mapping[f"rev{i}_desc"] = rev.description or ""

    # 填充剩餘的修訂記錄為空字串
    for i in range(len(schema.revision_records) + 1, 6):
        mapping[f"rev{i}_item"] = ""
        mapping[f"rev{i}_date"] = ""
        mapping[f"rev{i}_report_no"] = ""
        mapping[f"rev{i}_desc"] = ""

    # 附件
    if schema.attachments:
        mapping["attachment_list"] = "; ".join(schema.attachments)
    else:
        mapping["attachment_list"] = ""

    # ===================
    # Lab Fixed Info (實驗室固定資訊)
    # ===================
    mapping["lab_name"] = settings.lab_name
    mapping["lab_address"] = settings.lab_address
    mapping["lab_accreditation_no"] = settings.lab_accreditation_no
    mapping["lab_altitude"] = settings.lab_altitude

    # 報告預設值（如果 Schema 中沒有值則使用預設）
    if not mapping.get("test_type"):
        mapping["test_type"] = settings.default_test_type
    if not mapping.get("cns_standard"):
        mapping["cns_standard"] = settings.default_cns_standard
    if not mapping.get("cns_standard_version"):
        mapping["cns_standard_version"] = settings.default_cns_standard_version

    # ===================
    # Metadata
    # ===================
    mapping["extraction_timestamp"] = schema.extraction_timestamp or ""
    mapping["source_filename"] = schema.source_filename or ""

    logger.info(f"建立 placeholder mapping，共 {len(mapping)} 個欄位")
    return mapping


def build_checkbox_mapping(schema: ReportSchema) -> Dict[str, bool]:
    """
    根據 ReportSchema 建立 checkbox 狀態對應表

    這個 mapping 用於決定哪些 checkbox 需要打勾

    Args:
        schema: ReportSchema 物件

    Returns:
        Dict[checkbox_label, is_checked]
    """
    flags = schema.checkbox_flags
    mapping = {}

    # 產品群組
    mapping["AV"] = flags.is_av
    mapping["ICT"] = flags.is_ict
    mapping["Audio/Video & ICT"] = flags.is_av_ict
    mapping["AV & ICT"] = flags.is_av_ict
    mapping["Telecom"] = flags.is_telecom

    # 使用分類
    mapping["Ordinary"] = flags.is_ordinary
    mapping["Skilled"] = flags.is_skilled
    mapping["Instructed"] = flags.is_instructed

    # 電源等級
    mapping["Class I"] = flags.is_class_i
    mapping["Class II"] = flags.is_class_ii
    mapping["Class III"] = flags.is_class_iii

    # 移動性 / 設備移動性
    mapping["direct_plugin"] = flags.is_direct_plugin
    mapping["is_direct_plugin"] = flags.is_direct_plugin
    mapping["stationary"] = flags.is_stationary
    mapping["is_stationary"] = flags.is_stationary
    mapping["building_in"] = flags.is_building_in
    mapping["is_building_in"] = flags.is_building_in
    mapping["wall_ceiling"] = flags.is_wall_ceiling
    mapping["is_wall_ceiling"] = flags.is_wall_ceiling
    mapping["rack_mounted"] = flags.is_rack_mounted
    mapping["is_rack_mounted"] = flags.is_rack_mounted
    mapping["portable"] = flags.is_portable
    mapping["is_portable"] = flags.is_portable
    mapping["fixed"] = flags.is_fixed
    mapping["is_fixed"] = flags.is_fixed

    # 舊的 key (相容性)
    mapping["Portable"] = flags.is_portable
    mapping["Stationary"] = flags.is_stationary
    mapping["Fixed"] = flags.is_fixed

    # 連接類型
    mapping["Pluggable Type A"] = flags.is_pluggable_a
    mapping["Pluggable Type B"] = flags.is_pluggable_b
    mapping["Permanently Connected"] = flags.is_permanently_connected

    return mapping


# ==============================================
# Run Text Handling (解決 Word 切割問題)
# ==============================================

def get_paragraph_text(paragraph: Paragraph) -> str:
    """取得段落的完整文字（合併所有 runs）"""
    return "".join(run.text for run in paragraph.runs)


def get_cell_text(cell: _Cell) -> str:
    """取得表格儲存格的完整文字"""
    return "\n".join(get_paragraph_text(p) for p in cell.paragraphs)


def replace_text_in_runs(runs: List[Run], old_text: str, new_text: str) -> bool:
    """
    在多個 runs 中替換文字，保持格式

    這個函式處理 Word 將單一文字切割成多個 runs 的情況。
    例如 "{{name}}" 可能被切成 ["{{", "name", "}}"]

    策略：
    1. 先合併所有 runs 的文字
    2. 找到 old_text 的位置
    3. 計算每個 run 對應的字元範圍
    4. 精確地替換，保持格式

    Args:
        runs: Run 物件列表
        old_text: 要替換的文字
        new_text: 替換後的文字

    Returns:
        是否有進行替換
    """
    if not runs:
        return False

    # 合併所有文字
    full_text = "".join(run.text for run in runs)

    # 檢查是否包含目標文字
    if old_text not in full_text:
        return False

    # 計算每個 run 的字元範圍
    run_ranges = []  # [(start, end, run), ...]
    current_pos = 0
    for run in runs:
        run_len = len(run.text)
        run_ranges.append((current_pos, current_pos + run_len, run))
        current_pos += run_len

    # 找到 old_text 的位置
    start_idx = full_text.find(old_text)
    end_idx = start_idx + len(old_text)

    # 建立新的文字（替換後）
    new_full_text = full_text[:start_idx] + new_text + full_text[end_idx:]

    # 重新分配文字到各個 runs
    # 策略：把所有文字放到第一個受影響的 run，清空其他受影響的 runs
    for run_start, run_end, run in run_ranges:
        if run_end <= start_idx:
            # 這個 run 在替換區域之前，保持不變
            continue
        elif run_start >= end_idx:
            # 這個 run 在替換區域之後
            # 需要調整位置（因為替換可能改變長度）
            offset = len(new_text) - len(old_text)
            new_run_start = run_start + offset
            new_run_end = run_end + offset
            run.text = new_full_text[new_run_start:new_run_end]
        elif run_start <= start_idx and run_end >= end_idx:
            # 整個 old_text 都在這個 run 中
            run.text = run.text.replace(old_text, new_text, 1)
        elif run_start <= start_idx < run_end:
            # old_text 從這個 run 開始
            # 把這個 run 之前的部分 + new_text + 後續所有文字都放到這個 run
            run.text = new_full_text[run_start:]
            # 清空後續的 runs
            for _, _, other_run in run_ranges:
                if other_run != run and run_ranges.index((_, _, other_run)) > run_ranges.index((run_start, run_end, run)):
                    other_run.text = ""
            break
        elif start_idx < run_start < end_idx:
            # 這個 run 在 old_text 中間，清空它
            run.text = ""

    return True


def replace_placeholder_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]) -> int:
    """
    替換段落中的所有 placeholder

    Args:
        paragraph: Paragraph 物件
        mapping: placeholder 對應表

    Returns:
        替換的數量
    """
    count = 0
    full_text = get_paragraph_text(paragraph)

    # 找到所有 placeholder
    placeholders = PLACEHOLDER_PATTERN.findall(full_text)

    for placeholder in placeholders:
        key = placeholder.strip()
        if key in mapping:
            old_text = "{{" + placeholder + "}}"
            new_text = mapping[key]

            if replace_text_in_runs(paragraph.runs, old_text, str(new_text)):
                count += 1
                logger.debug(f"替換 {old_text} → {new_text[:50]}...")

    return count


def replace_checkbox_in_paragraph(paragraph: Paragraph, checkbox_mapping: Dict[str, bool]) -> int:
    """
    處理段落中的 checkbox

    尋找格式為「□ Label」或「[Label] □」的 checkbox，
    如果對應的 label 在 mapping 中為 True，則將 □ 改為 ■

    Args:
        paragraph: Paragraph 物件
        checkbox_mapping: checkbox 狀態對應表

    Returns:
        處理的數量
    """
    count = 0
    full_text = get_paragraph_text(paragraph)

    for label, is_checked in checkbox_mapping.items():
        if not is_checked:
            continue

        # 檢查是否包含這個 label 和 checkbox
        # 支援多種格式：
        # - □ Label
        # - [Label] □
        # - Label □
        # - □Label

        patterns = [
            f"□ {label}",
            f"□{label}",
            f"[{label}] □",
            f"{label} □",
            f"□  {label}",  # 雙空格
        ]

        for pattern in patterns:
            if pattern in full_text:
                # 找到了，把 □ 改成 ■
                # 但要小心只改這個 label 相關的 □
                if replace_text_in_runs(paragraph.runs, CHECKBOX_UNCHECKED, CHECKBOX_CHECKED):
                    count += 1
                    logger.debug(f"勾選 checkbox: {label}")
                    break

    return count


# ==============================================
# Table Processing
# ==============================================

def process_table(table: Table, mapping: Dict[str, str], checkbox_mapping: Dict[str, bool]) -> int:
    """
    處理表格中的所有儲存格

    Args:
        table: Table 物件
        mapping: placeholder 對應表
        checkbox_mapping: checkbox 狀態對應表

    Returns:
        替換的總數量
    """
    count = 0

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += replace_placeholder_in_paragraph(paragraph, mapping)
                count += replace_checkbox_in_paragraph(paragraph, checkbox_mapping)

    return count


# ==============================================
# Advanced Block/Table Helpers
# ==============================================

def replace_text_globally(doc: Document, pattern: str, replacement: str) -> int:
    """
    以簡單字串替換整份文件（含表格、頁首頁尾）中的文字。
    適用於固定值改成動態值（例如舊案編號、型號）。
    """
    count = 0
    def _replace_in_paragraph(paragraph: Paragraph) -> int:
        nonlocal count
        if not pattern or pattern not in paragraph.text:
            return 0
        for run in paragraph.runs:
            if pattern in run.text:
                run.text = run.text.replace(pattern, replacement)
        count += 1
        return 1

    def _replace_in_container(paragraphs):
        for p in paragraphs:
            _replace_in_paragraph(p)

    _replace_in_container(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_container(cell.paragraphs)
    for section in doc.sections:
        _replace_in_container(section.header.paragraphs)
        _replace_in_container(section.footer.paragraphs)
        for table in section.header.tables:
            for cell in table._cells:
                _replace_in_container(cell.paragraphs)
        for table in section.footer.tables:
            for cell in table._cells:
                _replace_in_container(cell.paragraphs)
    return count


def find_table_by_text_or_placeholder(doc: Document, keyword: str, placeholder: Optional[str] = None) -> Optional[Table]:
    """
    在文件與頁首頁尾搜尋含有關鍵字或 placeholder 的表格。
    keyword: 用於舊模板中的中文標記
    placeholder: 用於新模板中的 {{...}} 標記
    """
    def _search(tables):
        for table in tables:
            for cell in table._cells:
                txt = cell.text
                if (keyword and keyword in txt) or (placeholder and placeholder in txt):
                    return table
        return None

    table = _search(doc.tables)
    if table:
        return table
    for section in doc.sections:
        table = _search(section.header.tables)
        if table:
            return table
        table = _search(section.footer.tables)
        if table:
            return table
    return None


def fill_table_with_rows(table: Table, headers: List[str], rows: List[List[Any]]) -> None:
    """
    以指定的表頭與資料行覆寫現有表格內容。
    - 會新增或刪除行以符合 rows 數量
    - 會覆寫前 len(headers) 個欄位文字，其餘清空
    """
    if table is None:
        return

    required_rows = 1 + len(rows)
    # 新增行至足夠
    while len(table.rows) < required_rows:
        table.add_row()
    # 多餘行刪除
    while len(table.rows) > required_rows:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)

    # 表頭
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = str(headers[idx]) if idx < len(headers) else ""
    # 資料行
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(cells):
            cells[c_idx].text = str(row[c_idx]) if c_idx < len(row) and row[c_idx] is not None else ""


def render_factory_list(doc: Document, factories: List['FactoryInfo']) -> None:
    """將工廠清單填入封面表格的生產廠場欄位（若存在）。"""
    if not factories:
        return
    summary = "; ".join(" / ".join(filter(None, [f.name, f.address])) for f in factories if f.name or f.address)
    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if "生產廠場" in cell.text and idx + 1 < len(row.cells):
                    row.cells[idx + 1].text = summary


def insert_table_at_placeholder(doc: Document, placeholder: str, headers: List[str], rows: List[List[Any]]) -> bool:
    """
    尋找包含 placeholder 的段落或表格儲存格，插入表格並移除 placeholder。
    回傳是否插入成功。
    """
    def _create_table():
        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        for ci, cell in enumerate(tbl.rows[0].cells):
            cell.text = str(headers[ci]) if ci < len(headers) else ""
        for ri, data in enumerate(rows, start=1):
            for ci, cell in enumerate(tbl.rows[ri].cells):
                cell.text = str(data[ci]) if ci < len(data) else ""
        return tbl

    # 段落尋找
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            _create_table()
            paragraph.text = ""
            return True

    # 表格儲存格尋找
    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    _create_table()
                    paragraph.text = ""
                    return True
    return False


def render_factory_table_block(doc: Document, factories: List['FactoryInfo']) -> None:
    if not factories:
        return
    headers = ["Factory name", "Address"]
    rows = [[f.name, f.address] for f in factories]
    if not insert_table_at_placeholder(doc, "{{#BLOCK:FACTORY_TABLE_BLOCK}}", headers, rows):
        para = doc.add_paragraph("工廠清單")
        tbl = doc.add_table(rows=len(rows) + 1, cols=2)
        tbl.rows[0].cells[0].text = headers[0]
        tbl.rows[0].cells[1].text = headers[1]
        for ri, data in enumerate(rows, start=1):
            tbl.rows[ri].cells[0].text = str(data[0])
            tbl.rows[ri].cells[1].text = str(data[1])
    replace_text_globally(doc, "{{#BLOCK:FACTORY_TABLE_BLOCK}}", "")


def render_input_test_table(doc: Document, key_tables: 'KeyTables') -> None:
    """用 key_tables.input_test_raw 或 input_tests 填寫 B.2.5 表格。"""
    table = find_table_by_text_or_placeholder(doc, "表格: 輸入試驗", "{{#BLOCK:TABLE_B2_5_INPUT_TEST}}")
    # 資料來源優先 raw
    rows: List[List[Any]] = []
    if key_tables.input_test_raw:
        headers = key_tables.input_test_raw[0]
        rows = key_tables.input_test_raw[1:]
    elif key_tables.input_tests:
        headers = ["Voltage", "Frequency", "Current", "Power", "Condition", "Remarks"]
        for row in key_tables.input_tests:
            rows.append([
                row.voltage or "",
                row.frequency or "",
                row.current or "",
                row.power or "",
                row.test_condition or "",
                row.remarks or ""
        ])
    else:
        return
    if rows and table:
        fill_table_with_rows(table, headers, rows)
    elif rows:
        insert_table_at_placeholder(doc, "{{#BLOCK:TABLE_B2_5_INPUT_TEST}}", headers, rows)


def render_abnormal_fault_table(doc: Document, key_tables: 'KeyTables') -> None:
    """用 key_tables.abnormal_fault_raw 填寫 B.3/B.4 表格。"""
    table = find_table_by_text_or_placeholder(doc, "異常操作和故障條件試驗", "{{#BLOCK:TABLE_B3_B4_ABNORMAL_FAULT}}")
    if not key_tables.abnormal_fault_raw:
        return
    headers = key_tables.abnormal_fault_raw[0]
    rows = key_tables.abnormal_fault_raw[1:]
    if table:
        fill_table_with_rows(table, headers, rows)
    else:
        insert_table_at_placeholder(doc, "{{#BLOCK:TABLE_B3_B4_ABNORMAL_FAULT}}", headers, rows)
    # 若關鍵值因合併儲存格而未顯示，將所有行的合併文字附加在首格
    if table:
        table_text = " ".join(c.text for c in table._cells)
        combined = "\n".join(" ; ".join(str(v) for v in r if v) for r in rows)
        if any((str(v) and str(v) not in table_text) for r in rows for v in r):
            cell = table.rows[0].cells[0]
            cell.text = (cell.text + "\n" + combined).strip() if cell.text else combined


def render_temperature_block(doc: Document, schema: 'ReportSchema') -> None:
    """
    將溫度/負載條件敘述插入 {{#BLOCK:TEMPERATURE_REQUIREMENTS_TABLE}}。
    優先使用 schema.basic_info.temperature_requirements_text，其次使用 key_tables.temperature_rise 製作摘要。
    """
    placeholder = "{{#BLOCK:TEMPERATURE_REQUIREMENTS_TABLE}}"
    text = ""
    if getattr(schema.basic_info, "temperature_requirements_text", None):
        text = schema.basic_info.temperature_requirements_text
    elif schema.key_tables and schema.key_tables.temperature_rise:
        # 簡易摘要：列出量測位置與溫升
        parts = []
        for row in schema.key_tables.temperature_rise[:5]:
            seg = " / ".join(filter(None, [row.location, f"{row.measured_temp}C" if row.measured_temp else None, f"rise {row.temp_rise}K" if row.temp_rise else None]))
            if seg:
                parts.append(seg)
        if parts:
            text = "Temperature rise summary: " + "; ".join(parts)

    if not text:
        replace_text_globally(doc, placeholder, "")
        return

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = text
            return
    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = text
                    return
    replace_text_globally(doc, placeholder, text)


def render_max_output_block(doc: Document, basic_info: 'BasicInfo') -> None:
    """插入最大輸出功率/電壓/電流敘述，替換 {{#BLOCK:MAX_OUTPUT_POWER_BLOCK}}。"""
    placeholder = "{{#BLOCK:MAX_OUTPUT_POWER_BLOCK}}"
    max_w = basic_info.max_output_w or ""
    max_v = basic_info.max_output_v or ""
    max_a = basic_info.max_output_a or ""
    text = ""
    if max_w or max_v or max_a:
        parts = []
        if max_w:
            parts.append(f"{max_w}W")
        if max_v and max_a:
            parts.append(f"{max_v}V, {max_a}A")
        elif max_v:
            parts.append(f"{max_v}V")
        elif max_a:
            parts.append(f"{max_a}A")
        text = "最大連續輸出功率為 " + " / ".join(parts)
    if not text and basic_info.rated_output_lines:
        # 從額定輸出行推估最大功率
        numbers = []
        for line in basic_info.rated_output_lines:
            match = re.findall(r"([0-9]+(?:\\.[0-9]+)?)\\s*W", line, flags=re.IGNORECASE)
            numbers.extend(match)
        if numbers:
            text = f"最大連續輸出功率為 {max(numbers, key=lambda x: float(x))}W"

    if not text:
        replace_text_globally(doc, placeholder, "")
        return

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = text
            return
    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = text
                    return
    replace_text_globally(doc, placeholder, text)


def render_attachment_block(doc: Document, attachments: Optional[List[str]]) -> None:
    if not attachments:
        return
    placeholder = "{{#BLOCK:ATTACHMENT_LIST_BLOCK}}"

    def _write_list(paragraph):
        paragraph.text = ""
        for item in attachments:
            run = paragraph.add_run(f"• {item}\n")
        return True

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            _write_list(paragraph)
            return
    for table in doc.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    _write_list(paragraph)
                    return


def replace_placeholder_in_cell_text(cell: _Cell, key: str, value: str) -> int:
    count = 0
    for paragraph in cell.paragraphs:
        count += replace_placeholder_in_paragraph(paragraph, {key: value})
    return count


def replace_placeholders_in_textboxes(docx_path: str, mapping: Dict[str, str]) -> None:
    """
    直接修改 docx XML 以處理 python-docx 無法觸及的 textbox（w:txbxContent）。
    僅做簡單的字串替換，不改變結構。
    """
    if not mapping:
        return

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        word_dir = os.path.join(temp_dir, "word")
        targets = [os.path.join(word_dir, "document.xml")]
        targets += glob.glob(os.path.join(word_dir, "header*.xml"))
        targets += glob.glob(os.path.join(word_dir, "footer*.xml"))

        changed = False
        for xml_file in targets:
            if not os.path.exists(xml_file):
                continue
            with open(xml_file, "r", encoding="utf-8") as f:
                xml_text = f.read()
            original = xml_text
            for k, v in mapping.items():
                placeholder = f"{{{{{k}}}}}"
                if placeholder in xml_text:
                    xml_text = xml_text.replace(placeholder, v or "")
            if xml_text != original:
                with open(xml_file, "w", encoding="utf-8") as f:
                    f.write(xml_text)
                changed = True

        if changed:
            temp_docx = docx_path + ".txbx"
            with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        path = os.path.join(root, file)
                        arcname = os.path.relpath(path, temp_dir)
                        zipf.write(path, arcname)
            shutil.move(temp_docx, docx_path)
    except Exception as e:
        logger.warning(f"Textbox placeholder replacement failed: {e}")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


# ==============================================
# Main Fill Function
# ==============================================

def fill_cns_template(
    schema: ReportSchema,
    template_path: str,
    output_path: str,
    user_inputs: Optional[Dict[str, str]] = None
) -> None:
    """
    主要函式：填寫 CNS Word 模板

    這是此模組的主要入口點。

    設計原則：
    - 只修改文字內容，不改變文件結構
    - 使用 placeholder 替換機制（{{...}}）
    - 處理 checkbox 的勾選（□ → ■）
    - 保持原有格式（字體、大小、顏色、邊框等）

    Args:
        schema: ReportSchema 物件（包含所有要填寫的資料）
        template_path: Word 模板檔案路徑
        output_path: 輸出檔案路徑
        user_inputs: 使用者從前端輸入的額外欄位（選填）
            - report_author: 報告撰寫人
            - report_signer: 報告簽署人
            - series_model: 系列型號（逗號分隔）

    Raises:
        FileNotFoundError: 當模板檔案不存在時
        Exception: 當處理過程發生錯誤時

    Usage:
        >>> schema = ReportSchema(...)
        >>> fill_cns_template(schema, "templates/cns_template.docx", "output/report.docx")
    """
    logger.info(f"開始填寫 Word 模板: {template_path}")

    # 檢查模板檔案是否存在
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板檔案不存在: {template_path}")

    # 建立 mapping
    placeholder_mapping = build_placeholder_mapping(schema)
    checkbox_mapping = build_checkbox_mapping(schema)

    # 加入使用者輸入的欄位（如果有的話）
    if user_inputs:
        # 台灣申請者資訊（覆蓋 CB 報告中的製造商資訊）
        if user_inputs.get("applicant_name"):
            placeholder_mapping["applicant_name"] = user_inputs["applicant_name"]
            placeholder_mapping["applicant_zh"] = user_inputs["applicant_name"]
            placeholder_mapping["applicant_en"] = user_inputs["applicant_name"]
            logger.info(f"使用者輸入 - 申請者名稱: {user_inputs['applicant_name']}")
        if user_inputs.get("applicant_address"):
            placeholder_mapping["applicant_address"] = user_inputs["applicant_address"]
            placeholder_mapping["applicant_address_zh"] = user_inputs["applicant_address"]
            placeholder_mapping["applicant_address_en"] = user_inputs["applicant_address"]
            logger.info(f"使用者輸入 - 申請者地址: {user_inputs['applicant_address']}")
        if user_inputs.get("cns_report_no"):
            placeholder_mapping["report_no"] = user_inputs["cns_report_no"]
            placeholder_mapping["cns_report_no"] = user_inputs["cns_report_no"]
            logger.info(f"使用者輸入 - CNS 報告編號: {user_inputs['cns_report_no']}")

        # 其他欄位
        if user_inputs.get("report_author"):
            placeholder_mapping["report_author"] = user_inputs["report_author"]
            logger.info(f"使用者輸入 - 報告撰寫人: {user_inputs['report_author']}")
        if user_inputs.get("report_signer"):
            placeholder_mapping["report_signer"] = user_inputs["report_signer"]
            logger.info(f"使用者輸入 - 報告簽署人: {user_inputs['report_signer']}")
        if user_inputs.get("series_model"):
            placeholder_mapping["series_model"] = user_inputs["series_model"]
            logger.info(f"使用者輸入 - 系列型號: {user_inputs['series_model']}")

    # 載入模板
    doc = Document(template_path)

    total_replacements = 0

    # 處理文件主體的段落
    logger.info("處理文件主體...")
    for paragraph in doc.paragraphs:
        total_replacements += replace_placeholder_in_paragraph(paragraph, placeholder_mapping)
        total_replacements += replace_checkbox_in_paragraph(paragraph, checkbox_mapping)

    # 處理表格
    logger.info(f"處理 {len(doc.tables)} 個表格...")
    for table in doc.tables:
        total_replacements += process_table(table, placeholder_mapping, checkbox_mapping)

    # 處理頁首
    logger.info("處理頁首...")
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            total_replacements += replace_placeholder_in_paragraph(paragraph, placeholder_mapping)

        # 處理頁首中的表格
        for table in header.tables:
            total_replacements += process_table(table, placeholder_mapping, checkbox_mapping)

    # 處理頁尾
    logger.info("處理頁尾...")
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            total_replacements += replace_placeholder_in_paragraph(paragraph, placeholder_mapping)

        # 處理頁尾中的表格
        for table in footer.tables:
            total_replacements += process_table(table, placeholder_mapping, checkbox_mapping)

    # 補充：工廠清單、動態表格、型號/編號替換
    render_factory_list(doc, schema.factories)
    render_factory_table_block(doc, schema.factories)
    render_input_test_table(doc, schema.key_tables)
    render_abnormal_fault_table(doc, schema.key_tables)
    render_temperature_block(doc, schema)
    render_max_output_block(doc, schema.basic_info)
    render_attachment_block(doc, schema.attachments)

    main_model_for_cleanup = placeholder_mapping.get("main_model") or placeholder_mapping.get("model_main") or ""
    if main_model_for_cleanup:
        replace_text_globally(doc, "MC-601", main_model_for_cleanup)
    if placeholder_mapping.get("header_report_no"):
        replace_text_globally(doc, "AST-B-25120522-000", placeholder_mapping["header_report_no"])
    # 清理已知 placeholder / 舊案號
    replace_text_globally(doc, "{{TABLE_B2_5_INPUT_TEST}}", "")
    replace_text_globally(doc, "{{TABLE_B3_B4_ABNORMAL_FAULT}}", "")
    replace_text_globally(doc, "{{#BLOCK:TABLE_B2_5_INPUT_TEST}}", "")
    replace_text_globally(doc, "{{#BLOCK:TABLE_B3_B4_ABNORMAL_FAULT}}", "")
    replace_text_globally(doc, "{{report_author}}", "")
    replace_text_globally(doc, "{{report_signer}}", "")
    replace_text_globally(doc, "{{FACTORY_TABLE_BLOCK}}", "")
    replace_text_globally(doc, "{{ATTACHMENT_LIST_BLOCK}}", "")
    replace_text_globally(doc, "{{#BLOCK:FACTORY_TABLE_BLOCK}}", "")
    replace_text_globally(doc, "{{#BLOCK:ATTACHMENT_LIST_BLOCK}}", "")
    replace_text_globally(doc, "2025112058855971-00", "")
    replace_text_globally(doc, "DK-174052-UL", "")

    # 確保輸出目錄存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 儲存結果
    doc.save(output_path)

    logger.info(f"Word 模板填寫完成，共替換 {total_replacements} 處")
    logger.info(f"輸出檔案: {output_path}")

    # 更新 FORMCHECKBOX 狀態
    # 由於 python-docx 不支援 FORMCHECKBOX，需要在儲存後另外處理
    # 先處理 textbox 內殘留的 placeholder（python-docx 無法觸及）
    try:
        replace_placeholders_in_textboxes(output_path, placeholder_mapping)
    except Exception as e:
        logger.warning(f"Textbox 佔位符替換失敗: {e}")

    # 更新 FORMCHECKBOX 狀態
    # 由於 python-docx 不支援 FORMCHECKBOX，需要在儲存後另外處理
    try:
        update_formcheckbox_in_docx(
            output_path,
            schema.checkbox_flags,
            schema.test_item_particulars
        )
    except Exception as e:
        logger.warning(f"更新 FORMCHECKBOX 時發生錯誤（不影響其他功能）: {e}")

    # 基本檢核
    try:
        expected_models = {schema.basic_info.model_main} if schema.basic_info.model_main else set()
        expected_models.update([m.model for m in schema.series_models if m.model])
        expected_report_nos = set(filter(None, [
            placeholder_mapping.get("header_report_no"),
            placeholder_mapping.get("report_no"),
            placeholder_mapping.get("cb_report_no"),
            placeholder_mapping.get("cns_report_no"),
            schema.basic_info.ast_report_no,
            schema.basic_info.cb_report_no,
            schema.basic_info.cns_report_no
        ]))
        expected_factories = [f.name for f in schema.factories if f.name]
        expected_max_w = None
        if schema.basic_info.max_output_w:
            try:
                expected_max_w = float(str(schema.basic_info.max_output_w).replace("W","").strip())
            except Exception:
                expected_max_w = None
        if not expected_max_w and placeholder_mapping.get("max_output_w"):
            try:
                expected_max_w = float(str(placeholder_mapping.get("max_output_w")))
            except Exception:
                expected_max_w = None
        post_render_validate(
            output_path,
            {
                "ast_report_no": placeholder_mapping.get("header_report_no"),
                "models": list(expected_models),
                "report_numbers": list(expected_report_nos),
                "max_output_w": expected_max_w,
                "test_date_range": (schema.basic_info.test_date_from, schema.basic_info.test_date_to),
                "factories": expected_factories
            }
        )
    except Exception as e:
        logger.warning(f"產出檢核失敗: {e}")


# ==============================================
# Advanced Placeholder Functions
# ==============================================

def find_unreplaced_placeholders(doc_path: str) -> List[str]:
    """
    找出文件中尚未被替換的 placeholder

    用於 debug 或驗證模板

    Args:
        doc_path: Word 文件路徑

    Returns:
        未替換的 placeholder 列表
    """
    doc = Document(doc_path)
    unreplaced = set()

    # 檢查段落
    for paragraph in doc.paragraphs:
        text = get_paragraph_text(paragraph)
        matches = PLACEHOLDER_PATTERN.findall(text)
        unreplaced.update(matches)

    # 檢查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = get_cell_text(cell)
                matches = PLACEHOLDER_PATTERN.findall(text)
                unreplaced.update(matches)

    # 檢查頁首頁尾
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = get_paragraph_text(paragraph)
            matches = PLACEHOLDER_PATTERN.findall(text)
            unreplaced.update(matches)

        for paragraph in section.footer.paragraphs:
            text = get_paragraph_text(paragraph)
            matches = PLACEHOLDER_PATTERN.findall(text)
            unreplaced.update(matches)

    return sorted(list(unreplaced))


def extract_all_text(doc: Document) -> str:
    """彙整文件內所有文字（段落、表格、頁首頁尾）。"""
    texts: List[str] = []
    texts.extend(p.text for p in doc.paragraphs if p.text)
    for table in doc.tables:
        for cell in table._cells:
            if cell.text:
                texts.append(cell.text)
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs if p.text)
        texts.extend(p.text for p in section.footer.paragraphs if p.text)
        for tbl in section.header.tables:
            texts.extend(c.text for c in tbl._cells if c.text)
        for tbl in section.footer.tables:
            texts.extend(c.text for c in tbl._cells if c.text)
    return "\n".join(texts)


def post_render_validate(doc_path: str, expected: Dict[str, Any]) -> None:
    """
    基本產出檢核：
    - 禁止舊案型號/報告號
    - 頁首/頁尾報告編號一致
    - 危險 placeholder 殘留
    - 最大功率/日期/工廠/型號差異
    """
    doc = Document(doc_path)
    text = extract_all_text(doc)
    errors: List[str] = []

    if re.search(r"DK-\d+|2025112058855971-00|25120522-000", text):
        errors.append("Legacy CB/AST report number residual detected")
    if PLACEHOLDER_REGEX_RAW.search(text):
        errors.append("Unreplaced placeholder detected")

    expected_report_no = expected.get("ast_report_no") or expected.get("report_no")
    allow_reports = set(expected.get("report_numbers") or [])
    if expected_report_no:
        allow_reports.add(expected_report_no)

    if expected_report_no:
        header_ok = False
        footer_checked = False
        for section in doc.sections:
            for p in section.header.paragraphs:
                if expected_report_no in (p.text or ""):
                    header_ok = True
            for p in section.footer.paragraphs:
                if p.text:
                    footer_checked = footer_checked or ("AST-B" in p.text or expected_report_no in p.text)
                    if expected_report_no in p.text:
                        footer_checked = True
        if not header_ok:
            errors.append("Header report number mismatch")
        if footer_checked:
            footer_ok = any(expected_report_no in (p.text or "") for section in doc.sections for p in section.footer.paragraphs)
            if not footer_ok:
                errors.append("Footer report number mismatch")

    # 型號檢核：若提供 expected models，則限制文中出現的型號需在集合內
    expected_models = set(expected.get("models") or [])
    if expected_models:
        for m in re.finditer(r"\b[A-Z0-9][A-Z0-9-]{2,}\b", text):
            token = m.group(0)
            if len(token) < 4:
                continue
            window = text[max(0, m.start()-10):m.end()+10]
            if any(k in window for k in ["型號","Model","系列","model","MODEL"]) and token not in expected_models:
                errors.append(f"Unexpected model detected: {token}")

    # AST 報告號碼檢核：出現的 AST-B-* 需在 allowlist
    if allow_reports:
        ast_numbers = set(re.findall(r"AST-B-[0-9A-Za-z_-]+", text))
        unexpected_ast = {n for n in ast_numbers if n not in allow_reports}
        if unexpected_ast:
            errors.append(f"Unexpected AST report numbers: {', '.join(sorted(unexpected_ast))}")

    # Report number tokens (CB/AST) 檢核
    if allow_reports:
        report_tokens = set(re.findall(r"\b[0-9]{8,}-[0-9]{2}\b", text))
        unexpected_reports = {n for n in report_tokens if n not in allow_reports}
        if unexpected_reports:
            errors.append(f"Unexpected report numbers: {', '.join(sorted(unexpected_reports))}")

    # 最大功率檢核：若提供 expected max_output_w，檢查包含 MAX/最大 文字的 W 數值
    max_w = expected.get("max_output_w")
    if max_w:
        for m in re.finditer(r"([0-9]+(?:\.[0-9]+)?)\s*W", text, flags=re.IGNORECASE):
            val = float(m.group(1))
            ctx = text[max(0, m.start()-10):m.end()+10]
            if any(k in ctx for k in ["max", "MAX", "最大", "連續"]) and abs(val - float(max_w)) > 0.01:
                errors.append(f"Unexpected max power value: {val}W (expected {max_w}W)")

    # 日期區間檢核
    date_range = expected.get("test_date_range") or ()
    if len(date_range) == 2 and all(date_range):
        start, end = date_range
        if start in text and end not in text:
            errors.append("Test date range missing end date")

    # 工廠檢核：名稱需全部出現
    factories = expected.get("factories") or []
    for name in factories:
        if name and name not in text:
            errors.append(f"Factory missing in output: {name}")

    # 型號差異段落不得為僅命名不同
    if re.search(r"僅.*命名.*不同|only.*name", text, flags=re.IGNORECASE):
        errors.append("Model differences text too weak (naming only)")

    # 空白句結構
    if "為，" in text or "為 , " in text:
        errors.append("Found dangling sentence with empty value (\"為，\")")

    if errors:
        raise ValueError("Post render validation failed: " + "; ".join(errors))


def list_all_placeholders(template_path: str) -> List[str]:
    """
    列出模板中所有的 placeholder

    用於了解模板需要哪些欄位

    Args:
        template_path: 模板檔案路徑

    Returns:
        所有 placeholder 名稱的列表（去重後）
    """
    return find_unreplaced_placeholders(template_path)


# ==============================================
# Utility Functions
# ==============================================

def validate_template(template_path: str, schema: ReportSchema) -> Dict[str, List[str]]:
    """
    驗證模板與 schema 的相容性

    Args:
        template_path: 模板檔案路徑
        schema: ReportSchema 物件

    Returns:
        驗證結果，包含：
        - matched: 有對應值的 placeholder
        - unmatched: 沒有對應值的 placeholder
        - unused: schema 中有但模板中沒有的欄位
    """
    template_placeholders = set(list_all_placeholders(template_path))
    mapping = build_placeholder_mapping(schema)
    mapping_keys = set(mapping.keys())

    return {
        "matched": sorted(list(template_placeholders & mapping_keys)),
        "unmatched": sorted(list(template_placeholders - mapping_keys)),
        "unused": sorted(list(mapping_keys - template_placeholders))
    }


def create_sample_template_content() -> str:
    """
    產生範例模板內容說明

    這個函式不會產生實際的 .docx 檔案，
    而是回傳說明文字，告訴使用者如何準備模板
    """
    return """
==============================================
CNS 報告 Word 模板準備指南
==============================================

請在 templates/ 資料夾中放置您的 CNS 報告 Word 模板（.docx 格式）。

## Placeholder 格式

在模板中，使用 {{placeholder_name}} 格式標記需要填寫的欄位。

### 基本資料欄位

{{report_no}}           - 報告編號
{{cb_report_no}}        - CB 報告編號
{{standard}}            - 適用標準
{{applicant_en}}        - 申請人（英文）
{{applicant_zh}}        - 申請人（中文）
{{applicant_address_en}} - 申請人地址（英文）
{{applicant_address_zh}} - 申請人地址（中文）
{{manufacturer_en}}     - 製造商（英文）
{{manufacturer_zh}}     - 製造商（中文）
{{product_name_en}}     - 產品名稱（英文）
{{product_name_zh}}     - 產品名稱（中文）
{{model_main}}          - 主型號
{{ratings_input}}       - 輸入額定值
{{ratings_output}}      - 輸出額定值
{{issue_date}}          - 發行日期

### 試驗樣品特性欄位

{{product_group}}       - 產品群組
{{ovc}}                 - 過電壓類別
{{pollution_degree}}    - 污染等級
{{ip_code}}             - IP 防護等級
{{tma}}                 - 最高環境溫度
{{altitude_limit_m}}    - 海拔高度限制

### 系列型號欄位（1-60）

{{series_model_1}}          - 第 1 個型號
{{series_model_1_vout}}     - 第 1 個型號輸出電壓
{{series_model_1_iout}}     - 第 1 個型號輸出電流
{{series_model_1_pout}}     - 第 1 個型號輸出功率
{{series_model_1_case_type}} - 第 1 個型號外殼類型

（以此類推到 {{series_model_60}}）

## Checkbox 格式

在模板中使用 □ 符號表示未勾選的 checkbox：

□ AV
□ ICT
□ Audio/Video & ICT
□ Ordinary
□ Skilled
□ Instructed
□ Class I
□ Class II
□ Class III

程式會根據報告內容自動將對應的 □ 改成 ■

## 注意事項

1. 請勿改變模板的段落、表格結構
2. Placeholder 可能被 Word 切成多個 run，程式會自動處理
3. 建議使用簡單的字型（如 Arial、新細明體）
4. 儲存時使用 .docx 格式
"""


# ==============================================
# Testing Functions
# ==============================================

def test_fill_with_mock_data(template_path: str, output_path: str) -> None:
    """
    使用模擬資料測試模板填寫

    Args:
        template_path: 模板路徑
        output_path: 輸出路徑
    """
    from services.azure_llm import create_mock_schema

    schema = create_mock_schema()
    fill_cns_template(schema, template_path, output_path)
    logger.info(f"測試完成，輸出: {output_path}")


# ==============================================
# FORMCHECKBOX XML Processing
# ==============================================

# FORMCHECKBOX 標籤對應表
# 格式: { "XML中的標籤文字": "checkbox_flags 中的屬性名" }
FORMCHECKBOX_LABEL_MAPPING = {
    # 產品群組
    "終端產品": None,  # 根據 is_av, is_ict 等判斷
    "內建元件": None,  # 根據產品類型判斷

    # 使用分類
    "普通": "is_ordinary",
    "普通人員": "is_ordinary",
    "兒童可能出現": None,  # 特殊處理
    "受指導人員": "is_instructed",
    "技術人員": "is_skilled",

    # 電源連接
    "AC mains": None,  # 根據 mains_supply 判斷
    "DC mains": None,

    # 電源等級
    "Class I": "is_class_i",
    "Class II": "is_class_ii",
    "Class III": "is_class_iii",

    # 移動性 / 設備移動性
    "直插式設備": "is_direct_plugin",
    "放置式設備": "is_stationary",
    "崁入式設備": "is_building_in",
    "壁面/天花板安裝式": "is_wall_ceiling",
    "SRME/機架安裝": "is_rack_mounted",
    "移動式設備": "is_portable",
    "手持式設備": "is_portable",
    "可攜式設備": "is_portable",

    # 污染等級 - 需要特殊處理
    "PD 1": None,
    "PD 2": None,
    "PD 3": None,

    # 電力系統
    "TN": None,
    "TT": None,
}


def update_formcheckbox_in_xml(
    xml_content: str,
    checkbox_flags: 'CheckboxFlags',
    test_item_particulars: 'TestItemParticulars'
) -> str:
    """
    更新 Word 文件 XML 中的 FORMCHECKBOX 狀態

    由於 python-docx 不支援操作 FORMCHECKBOX，我們需要直接修改 XML。

    Args:
        xml_content: document.xml 的內容
        checkbox_flags: CheckboxFlags 物件
        test_item_particulars: TestItemParticulars 物件

    Returns:
        更新後的 XML 內容
    """
    # 建立要勾選的標籤集合
    labels_to_check = set()
    labels_to_uncheck = set()

    # ===================
    # 使用分類（根據 checkbox_flags）
    # ===================
    # 模板中的標籤是「普通」而不是「普通人員」
    if checkbox_flags.is_ordinary:
        labels_to_check.add("普通")
    else:
        labels_to_uncheck.add("普通")

    if checkbox_flags.is_skilled:
        labels_to_check.add("技術人員")
    else:
        labels_to_uncheck.add("技術人員")

    if checkbox_flags.is_instructed:
        labels_to_check.add("受指導人員")
    else:
        labels_to_uncheck.add("受指導人員")

    # ===================
    # 移動性 / 設備移動性
    # ===================
    # 直插式設備
    if checkbox_flags.is_direct_plugin:
        labels_to_check.add("直插式設備")
    else:
        labels_to_uncheck.add("直插式設備")

    # 放置式設備
    if checkbox_flags.is_stationary:
        labels_to_check.add("放置式設備")
    else:
        labels_to_uncheck.add("放置式設備")

    # 崁入式設備
    if checkbox_flags.is_building_in:
        labels_to_check.add("崁入式設備")
    else:
        labels_to_uncheck.add("崁入式設備")

    # 壁面/天花板安裝式
    if checkbox_flags.is_wall_ceiling:
        labels_to_check.add("壁面/天花板安裝式")
    else:
        labels_to_uncheck.add("壁面/天花板安裝式")

    # SRME/機架安裝
    if checkbox_flags.is_rack_mounted:
        labels_to_check.add("SRME/機架安裝")
    else:
        labels_to_uncheck.add("SRME/機架安裝")

    # 可攜式/移動式/手持式
    if checkbox_flags.is_portable:
        labels_to_check.add("移動式設備")
        labels_to_check.add("手持式設備")
        labels_to_check.add("可攜式設備")
    else:
        labels_to_uncheck.add("移動式設備")
        labels_to_uncheck.add("手持式設備")
        labels_to_uncheck.add("可攜式設備")

    # ===================
    # 根據 test_item_particulars 處理
    # ===================

    # 產品群組
    product_group = (test_item_particulars.product_group or "").upper()
    if "AV" in product_group or "ICT" in product_group or "終端" in product_group:
        labels_to_check.add("終端產品")
        labels_to_uncheck.add("內建元件")
    elif "COMPONENT" in product_group or "元件" in product_group:
        labels_to_check.add("內建元件")
        labels_to_uncheck.add("終端產品")

    # 電源類型
    mains_supply = (test_item_particulars.mains_supply or "").upper()
    if "AC" in mains_supply:
        labels_to_check.add("AC mains")
    else:
        labels_to_uncheck.add("AC mains")
    if "DC" in mains_supply:
        labels_to_check.add("DC mains")
    else:
        labels_to_uncheck.add("DC mains")
    # Not mains connected
    if "NOT" in mains_supply or "BATTERY" in mains_supply:
        labels_to_check.add("Not mains connected:")
    else:
        labels_to_uncheck.add("Not mains connected:")

    # 污染等級
    pollution = test_item_particulars.pollution_degree or ""
    if "1" in pollution:
        labels_to_check.add("PD 1")
        labels_to_uncheck.add("PD 2")
        labels_to_uncheck.add("PD 3")
    elif "2" in pollution:
        labels_to_check.add("PD 2")
        labels_to_uncheck.add("PD 1")
        labels_to_uncheck.add("PD 3")
    elif "3" in pollution:
        labels_to_check.add("PD 3")
        labels_to_uncheck.add("PD 1")
        labels_to_uncheck.add("PD 2")

    # 電力系統（通常是 TN）
    labels_to_check.add("TN")
    labels_to_uncheck.add("TT")
    labels_to_uncheck.add("IT -")

    logger.info(f"FORMCHECKBOX - 要勾選的標籤: {labels_to_check}")
    logger.info(f"FORMCHECKBOX - 要取消勾選的標籤: {labels_to_uncheck}")

    # 處理 XML
    # 策略：找到每個 FORMCHECKBOX，檢查它後面的標籤，決定是否勾選

    modified_xml = xml_content
    changes_made = 0

    # 找所有 w:tc（儲存格）區塊
    # 對於每個儲存格，找其中的 checkbox 和標籤對應關係

    def process_checkbox_cell(cell_xml: str) -> str:
        """處理單個儲存格中的 checkbox"""
        nonlocal changes_made

        # 找所有 checkbox 的位置
        checkbox_pattern = r'(<w:checkBox>)(.*?)(</w:checkBox>)'

        # 找所有文字標籤
        text_positions = []
        for m in re.finditer(r'<w:t[^>]*>([^<]+)</w:t>', cell_xml):
            text = m.group(1).strip()
            if text and text not in [' ', '\t', '\n']:
                text_positions.append((m.start(), text))

        # 處理每個 checkbox
        result = cell_xml
        offset = 0

        for cb_match in re.finditer(checkbox_pattern, cell_xml, re.DOTALL):
            cb_start = cb_match.start()
            cb_inner = cb_match.group(2)

            # 找這個 checkbox 後面最近的標籤
            label = None
            for text_pos, text in text_positions:
                if text_pos > cb_match.end():
                    label = text
                    break

            if not label:
                continue

            # 決定是否勾選
            should_check = label in labels_to_check
            should_uncheck = label in labels_to_uncheck

            if not should_check and not should_uncheck:
                continue

            # 檢查當前狀態
            is_currently_checked = '<w:checked/>' in cb_inner or '<w:checked w:val="1"/>' in cb_inner

            if should_check and not is_currently_checked:
                # 需要勾選：添加 <w:checked/>
                new_inner = cb_inner.rstrip()
                if not new_inner.endswith('<w:checked/>'):
                    # 在 </w:checkBox> 前插入 <w:checked/>
                    new_cb = f'<w:checkBox>{cb_inner}<w:checked/></w:checkBox>'
                    old_cb = cb_match.group(0)
                    result = result[:cb_match.start() + offset] + new_cb + result[cb_match.end() + offset:]
                    offset += len(new_cb) - len(old_cb)
                    changes_made += 1
                    logger.debug(f"FORMCHECKBOX - 勾選: {label}")

            elif should_uncheck and is_currently_checked:
                # 需要取消勾選：移除 <w:checked/>
                new_inner = re.sub(r'<w:checked[^>]*/>', '', cb_inner)
                new_cb = f'<w:checkBox>{new_inner}</w:checkBox>'
                old_cb = cb_match.group(0)
                result = result[:cb_match.start() + offset] + new_cb + result[cb_match.end() + offset:]
                offset += len(new_cb) - len(old_cb)
                changes_made += 1
                logger.debug(f"FORMCHECKBOX - 取消勾選: {label}")

        return result

    # 找「試驗樣品特性」表格並處理
    # 這個表格包含大部分需要更新的 checkbox
    table_match = re.search(r'(<w:tbl>.*?試驗樣品特性.*?</w:tbl>)', modified_xml, re.DOTALL)
    if table_match:
        table_xml = table_match.group(1)
        new_table_xml = process_checkbox_cell(table_xml)
        modified_xml = modified_xml[:table_match.start()] + new_table_xml + modified_xml[table_match.end():]

    logger.info(f"FORMCHECKBOX - 共修改 {changes_made} 個 checkbox")

    return modified_xml


def update_formcheckbox_in_docx(
    docx_path: str,
    checkbox_flags: 'CheckboxFlags',
    test_item_particulars: 'TestItemParticulars'
) -> None:
    """
    更新 .docx 檔案中的 FORMCHECKBOX 狀態

    .docx 檔案實際上是一個 ZIP 壓縮檔，包含多個 XML 檔案。
    此函式會：
    1. 解壓縮 .docx
    2. 修改 word/document.xml 中的 checkbox 狀態
    3. 重新壓縮為 .docx

    Args:
        docx_path: .docx 檔案路徑
        checkbox_flags: CheckboxFlags 物件
        test_item_particulars: TestItemParticulars 物件
    """
    import tempfile

    logger.info(f"開始更新 FORMCHECKBOX: {docx_path}")

    # 建立暫存目錄
    temp_dir = tempfile.mkdtemp()

    try:
        # 解壓縮 .docx
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # 讀取 document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()

        # 更新 checkbox
        new_xml_content = update_formcheckbox_in_xml(
            xml_content, checkbox_flags, test_item_particulars
        )

        # 寫回 document.xml
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(new_xml_content)

        # 重新壓縮為 .docx
        # 注意：需要保持原始的壓縮結構
        temp_docx = docx_path + '.tmp'
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)

        # 替換原檔案
        shutil.move(temp_docx, docx_path)

        logger.info(f"FORMCHECKBOX 更新完成: {docx_path}")

    finally:
        # 清理暫存目錄
        shutil.rmtree(temp_dir, ignore_errors=True)
